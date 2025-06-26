import streamlit as st
import os
import json
import re
import uuid
import extra_streamlit_components as esc
from typing import List, Dict, Any
from llama_index.core.llms import ChatMessage, MessageRole
from llama_index.core import Settings
import stui
from agent import create_orchestrator_agent, generate_suggested_prompts, SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS, initialize_settings as initialize_agent_settings, generate_llm_greeting
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import io # Import io module for BytesIO
from llama_index.core.tools import FunctionTool # Import FunctionTool

# Import necessary libraries for Hugging Face integration
from huggingface_hub import HfFileSystem 
import os # Import os to access environment variables

# Initialize HfFileSystem globally
fs = HfFileSystem() 
load_dotenv()

PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))

cookies = esc.CookieManager(key="esi_cookie_manager")

SIMPLE_STORE_PATH_RELATIVE = os.getenv("SIMPLE_STORE_PATH", "ragdb/simple_vector_store")
DB_PATH = os.path.join(PROJECT_ROOT, SIMPLE_STORE_PATH_RELATIVE)
AGENT_SESSION_KEY = "esi_orchestrator_agent"
DOWNLOAD_MARKER = "---DOWNLOAD_FILE---"
RAG_SOURCE_MARKER_PREFIX = "---RAG_SOURCE---"

# Import UI_ACCESSIBLE_WORKSPACE from tools.py
from tools import UI_ACCESSIBLE_WORKSPACE
# Import HF_USER_MEMORIES_DATASET_ID from config.py
from config import HF_USER_MEMORIES_DATASET_ID

# Constant to control the maximum number of messages sent in chat history to the LLM
MAX_CHAT_HISTORY_MESSAGES = 15 # Keep the last N messages to manage context length

@st.cache_resource
def setup_global_llm_settings() -> tuple[bool, str | None]:
    """Initializes global LLM settings using st.cache_resource to run only once."""
    print("Initializing LLM settings...")
    try:
        initialize_agent_settings()
        print("LLM settings initialized successfully.")
        return True, None
    except Exception as e:
        error_message = f"Fatal Error: Could not initialize LLM settings. {e}"
        print(error_message)
        return False, error_message

# New cached function for initial greeting
@st.cache_data(show_spinner=False)
def _get_initial_greeting_text():
    """Generates and caches the initial LLM greeting text for startup."""
    return generate_llm_greeting()

# New cached wrapper for suggested prompts
@st.cache_data(show_spinner=False)
def _cached_generate_suggested_prompts(chat_history: List[Dict[str, Any]]) -> List[str]:
    """
    Generates suggested prompts based on chat history, cached to avoid redundant LLM calls.
    The cache key is based on the content of chat_history.
    """
    print("Generating suggested prompts...")
    return generate_suggested_prompts(chat_history)

# Define dynamic tool functions that can access st.session_state
def read_uploaded_document_tool_fn(filename: str) -> str:
    """Reads the full text content of a document previously uploaded by the user.
    Input is the exact filename (e.g., 'my_dissertation.pdf')."""
    if "uploaded_documents" not in st.session_state or filename not in st.session_state.uploaded_documents:
        return f"Error: Document '{filename}' not found in uploaded documents. Available documents: {list(st.session_state.uploaded_documents.keys())}"
    return st.session_state.uploaded_documents[filename]

def analyze_dataframe_tool_fn(filename: str, head_rows: int = 5) -> str:
    """Provides summary information (shape, columns, dtypes, head, describe) about a pandas DataFrame
    previously uploaded by the user. Input is the exact filename (e.g., 'my_data.csv').
    For more complex analysis, use the 'code_interpreter' tool."""
    if "uploaded_dataframes" not in st.session_state or filename not in st.session_state.uploaded_dataframes:
        return f"Error: DataFrame '{filename}' not found in uploaded dataframes. Available dataframes: {list(st.session_state.uploaded_dataframes.keys())}"
    
    df = st.session_state.uploaded_dataframes[filename]
    
    info_str = f"DataFrame: {filename}\n"
    info_str += f"Shape: {df.shape}\n"
    info_str += f"Columns: {', '.join(df.columns)}\n"
    info_str += f"Data Types:\n{df.dtypes.to_string()}\n"
    
    # Ensure head_rows is not negative and not too large
    head_rows = max(0, min(head_rows, len(df)))
    if head_rows > 0:
        info_str += f"First {head_rows} rows:\n{df.head(head_rows).to_string()}\n"
    else:
        info_str += "No head rows requested or available.\n"

    info_str += f"Summary Statistics:\n{df.describe().to_string()}\n"
    
    return info_str

@st.cache_resource
def setup_agent(max_search_results: int) -> tuple[Any | None, str | None]:
    """Initializes the orchestrator agent using st.cache_resource to run only once per max_search_results value.
    Returns a tuple (agent_instance, error_message).
    agent_instance is None if an error occurred.
    error_message is None if successful.
    """
    print("Initializing AI agent...")
    try:
        # Create dynamic tools here, passing the functions defined above
        uploaded_doc_reader_tool = FunctionTool.from_defaults(
            fn=read_uploaded_document_tool_fn,
            name="read_uploaded_document",
            description="Reads the full text content of a document previously uploaded by the user. Input is the exact filename (e.g., 'my_dissertation.pdf'). Use this to answer questions about the content of uploaded documents."
        )

        dataframe_analyzer_tool = FunctionTool.from_defaults(
            fn=analyze_dataframe_tool_fn,
            name="analyze_uploaded_dataframe",
            description="Provides summary information (shape, columns, dtypes, head, describe) about a pandas DataFrame previously uploaded by the user. Input is the exact filename (e.g., 'my_data.csv'). Use this to understand the structure and basic statistics of uploaded datasets. For more complex analysis, use the 'code_interpreter' tool."
        )

        # Pass these dynamic tools and max_search_results to the agent creation function
        agent_instance = create_orchestrator_agent(
            dynamic_tools=[uploaded_doc_reader_tool, dataframe_analyzer_tool],
            max_search_results=max_search_results # Pass the parameter here
        )
        print("AI agent initialized successfully.")
        return agent_instance, None
    except Exception as e:
        error_message = f"Failed to initialize the AI agent. Please check configurations. Error: {e}"
        print(f"Error initializing AI agent: {e}")
        return None, error_message

def _get_or_create_user_id(long_term_memory_enabled_param: bool) -> tuple[str, str]:
    """
    Determines user ID and necessary cookie action.
    Returns a tuple: (user_id: str, cookie_action_flag: str).
    cookie_action_flag can be "DO_NOTHING", "SET_COOKIE", or "DELETE_COOKIE".
    This function NO LONGER performs cookie operations directly.
    """
    existing_user_id = cookies.get(cookie="user_id")

    if long_term_memory_enabled_param:
        if existing_user_id:
            return existing_user_id, "DO_NOTHING"
        else:
            new_user_id = str(uuid.uuid4())
            return new_user_id, "SET_COOKIE"
    else:  # Long-term memory is disabled
        temporary_user_id = str(uuid.uuid4())
        if existing_user_id:
            return temporary_user_id, "DELETE_COOKIE"
        else:
            return temporary_user_id, "DO_NOTHING"

@st.cache_resource
def _initialize_user_session_data(long_term_memory_enabled_param: bool) -> tuple[str, Dict[str, Any], Dict[str, Any], str]:
    """
    Initializes user ID, loads chat data from Hugging Face (if long-term memory is enabled),
    and returns the cookie action flag.
    This function is cached to run only once per Streamlit session, or when its parameters change.
    Returns: (user_id, chat_metadata, all_chat_messages, cookie_action_flag)
    """
    print("Initializing user session data...")

    user_id, cookie_action_flag = _get_or_create_user_id(long_term_memory_enabled_param)

    chat_metadata = {}
    all_chat_messages = {}

    if long_term_memory_enabled_param:
        print(f"Loading user data for user {user_id} from Hugging Face...")
        user_data = _load_user_data_from_hf(user_id) # This function is not cached, but its call is within a cached function
        chat_metadata = user_data["metadata"]
        all_chat_messages = user_data["messages"]
        print(f"Loaded {len(chat_metadata)} chats for user {user_id}.")
    else:
        print(f"Long-term memory disabled. No historical data loaded for temporary user_id {user_id}.")

    return user_id, chat_metadata, all_chat_messages, cookie_action_flag

def _load_user_data_from_hf(user_id: str) -> Dict[str, Any]:
    """
    Loads all chat metadata and histories for a user from JSON files on Hugging Face.
    This function is NOT cached by Streamlit.
    """
    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("HF_TOKEN environment variable not set. Cannot load user data from Hugging Face.")
        return {"metadata": {}, "messages": {}}

    all_chat_metadata = {}
    all_chat_messages = {}

    try:
        # Use HF_USER_MEMORIES_DATASET_ID for user memories
        metadata_filename_in_repo = f"user_memories/{user_id}_metadata.json"
        messages_filename_in_repo = f"user_memories/{user_id}_messages.json"

        # Construct the full HfFileSystem path
        metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{metadata_filename_in_repo}"
        messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"

        # Try to download and load metadata using HfFileSystem
        try:
            metadata_content = fs.read_text(metadata_hf_path, token=hf_token)
            all_chat_metadata = json.loads(metadata_content)
        except FileNotFoundError as e:
            print(f"Metadata file not found for user {user_id} at {metadata_hf_path}: {e}. Metadata will be empty.")
            all_chat_metadata = {}
        except Exception as e:
            print(f"Error loading metadata for user {user_id} from {metadata_hf_path}: {e}. Metadata will be empty.")
            all_chat_metadata = {}

        # Try to download and load messages using HfFileSystem
        try:
            messages_content = fs.read_text(messages_hf_path, token=hf_token)
            all_chat_messages = json.loads(messages_content)
        except FileNotFoundError as e:
            print(f"Messages file not found for user {user_id} at {messages_hf_path}: {e}. Messages will be empty.")
            all_chat_messages = {}
        except Exception as e:
            print(f"Error loading messages for user {user_id} from {messages_hf_path}: {e}. Messages will be empty.")
            all_chat_messages = {}

        return {"metadata": all_chat_metadata, "messages": all_chat_messages}

    except Exception as e:
        print(f"Error loading user data from Hugging Face for user {user_id}: {e}")
        return {"metadata": {}, "messages": {}}

def save_chat_history(user_id: str, chat_id: str, messages: List[Dict[str, Any]]):
    """
    Saves a specific chat history for a given user ID to a JSON file on Hugging Face.
    """
    if not st.session_state.long_term_memory_enabled:
        return

    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("HF_TOKEN environment variable not set. Skipping Hugging Face upload for chat history.")
        return

    try:
        # Use HF_USER_MEMORIES_DATASET_ID for user memories
        messages_filename_in_repo = f"user_memories/{user_id}_messages.json"
        messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"

        # Load existing messages, append the new chat, and save
        try:
            # Use fs.read_text to get the existing messages file content
            existing_messages_content = fs.read_text(messages_hf_path, token=hf_token)
            existing_messages = json.loads(existing_messages_content)
        except FileNotFoundError as e:
            print(f"Existing messages file not found at {messages_hf_path}: {e}. Starting with empty messages.")
            existing_messages = {}
        except Exception as e:
            print(f"Error loading existing messages from {messages_hf_path}: {e}. Starting with empty messages.")
            existing_messages = {}

        existing_messages[chat_id] = messages

        # Use fs.open to write the content
        with fs.open(messages_hf_path, "w", token=hf_token) as f:
            f.write(json.dumps(existing_messages, indent=2))
        
        print(f"Chat history for chat {chat_id} saved to {messages_filename_in_repo} on Hugging Face.")

    except Exception as e:
        print(f"Error saving chat history to Hugging Face for chat {chat_id} (user {user_id}): {e}")
        st.error(f"Error saving chat history to cloud: {e}")

def save_chat_metadata(user_id: str, chat_metadata: Dict[str, str]):
    """Saves the chat metadata (ID to name mapping) for a user to a JSON file on Hugging Face."""
    if not st.session_state.long_term_memory_enabled:
        return

    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("HF_TOKEN environment variable not set. Skipping Hugging Face upload for metadata.")
        return

    try:
        # Use HF_USER_MEMORIES_DATASET_ID for user memories
        metadata_filename_in_repo = f"user_memories/{user_id}_metadata.json"
        metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{metadata_filename_in_repo}"

        # Use fs.open to write the content
        with fs.open(metadata_hf_path, "w", token=hf_token) as f:
            f.write(json.dumps(chat_metadata, indent=2))
        
        print(f"Chat metadata for user {user_id} saved to {metadata_filename_in_repo} on Hugging Face.")

    except Exception as e:
        print(f"Error saving chat metadata to Hugging Face for user {user_id}: {e}")
        st.error(f"Error saving chat metadata to cloud: {e}")

def format_chat_history(streamlit_messages: List[Dict[str, Any]]) -> List[ChatMessage]:
    """
    Converts Streamlit message history to LlamaIndex ChatMessage list,
    truncating to the most recent messages to manage context length.
    """
    # Truncate messages to keep only the most recent ones
    # If the list is shorter than MAX_CHAT_HISTORY_MESSAGES, it will take all.
    truncated_messages = streamlit_messages[-MAX_CHAT_HISTORY_MESSAGES:]

    history = []
    for msg in truncated_messages:
        role = MessageRole.USER if msg["role"] == "user" else MessageRole.ASSISTANT
        history.append(ChatMessage(role=role, content=msg["content"]))
    return history

import time # Added for simulated streaming

def get_agent_response(query: str, chat_history: List[ChatMessage]): # -> Generator[str, None, None]:
    """
    Get a response from the agent stored in the session state using the chat method,
    explicitly passing the conversation history.
    This function simulates streaming by yielding words with a delay.
    """
    agent = st.session_state[AGENT_SESSION_KEY]

    try:
        # Get temperature from session state
        current_temperature = st.session_state.get("llm_temperature", 0.7)
        current_verbosity = st.session_state.get("llm_verbosity", 3) # Default to 3 if not found

        # New logic to set temperature using llama_index.core.Settings
        if Settings.llm:
            if hasattr(Settings.llm, 'temperature'):
                Settings.llm.temperature = current_temperature
            else:
                print(f"Warning: Settings.llm ({type(Settings.llm)}) does not have a 'temperature' attribute.")
        else:
            print("Warning: Settings.llm is not initialized. Cannot set temperature.")

        # Prepend verbosity level to the query
        modified_query = f"Verbosity Level: {current_verbosity}. {query}"

        # Simulate thinking before responding (optional, but can make spinner more meaningful)
        # time.sleep(0.5) # Removed as st.spinner is used by st.write_stream

        response = agent.chat(modified_query, chat_history=chat_history)
        response_text = response.response if hasattr(response, 'response') else str(response)

        words = response_text.split(" ")
        for word in words:
            yield word + " "
            time.sleep(0.05) # Small delay for simulated streaming effect

    except Exception as e:
        error_message = f"I apologize, but I encountered an error: {str(e)}"
        print(f"Error getting orchestrator agent response: {type(e).__name__} - {e}")
        yield error_message # Yield the error message as a single chunk

def create_new_chat_session_in_memory():
    """
    Creates a new chat session (ID, name, empty messages) in memory (st.session_state)
    and sets it as the current chat. Does NOT save to Hugging Face immediately.
    """
    new_chat_id = str(uuid.uuid4())
    
    new_chat_name = "Current Session" # Default for disabled memory
    if st.session_state.long_term_memory_enabled:
        existing_idea_nums = []
        for name in st.session_state.chat_metadata.values():
            match = re.match(r"Idea (\d+)", name)
            if match:
                existing_idea_nums.append(int(match.group(1)))
        
        next_idea_num = 1
        if existing_idea_nums:
            next_idea_num = max(existing_idea_nums) + 1
        new_chat_name = f"Idea {next_idea_num}"

    st.session_state.chat_metadata[new_chat_id] = new_chat_name
    # Keep generate_llm_greeting here for new chat creation
    st.session_state.all_chat_messages[new_chat_id] = [{"role": "assistant", "content": _get_initial_greeting_text()}]
    st.session_state.current_chat_id = new_chat_id
    st.session_state.messages = st.session_state.all_chat_messages[new_chat_id]
    st.session_state.chat_modified = False # New chats are initially unsaved
    
    # Generate initial prompts for the new chat
    st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages)

    print(f"Created new chat: '{new_chat_name}' (ID: {new_chat_id})")
    return new_chat_id # Return the new chat ID

def switch_chat(chat_id: str):
    """Switches to an existing chat, ensuring messages are loaded."""
    if not st.session_state.long_term_memory_enabled:
        print("Long-term memory disabled. Cannot switch to historical chats. Starting a new temporary session.")
        create_new_chat_session_in_memory()
        st.rerun() # Keep rerun here for user-initiated switch when LTM is off
        return

    if chat_id not in st.session_state.chat_metadata:
        print(f"Error: Attempted to switch to chat ID '{chat_id}' not found in metadata.")
        return

    # Messages for the target chat_id should already be loaded in st.session_state.all_chat_messages
    # by _initialize_user_session_data or _load_user_data_from_hf.
    # If for some reason they are not, it indicates an heinous issue with the loading logic.
    if st.session_state.all_chat_messages.get(chat_id) is None:
        print(f"WARNING: Messages for current chat ID '{chat_id}' were not loaded. Setting to empty list.")
        st.session_state.all_chat_messages[chat_id] = [] # Fallback
            
    st.session_state.messages = st.session_state.all_chat_messages.get(chat_id, [])
    st.session_state.current_chat_id = chat_id # Ensure current_chat_id is set here
    
    st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages) # Use cached version
    st.session_state.chat_modified = True # Assume existing chat is modified if switched to (will be saved on next AI response)
    print(f"Switched to chat: '{st.session_state.chat_metadata.get(chat_id, 'Unknown')}' (ID: {chat_id})")
    st.rerun() # Keep rerun here for user-initiated switch when LTM is on

def delete_chat_session(chat_id: str):
    """Deletes a chat history and its metadata from Hugging Face."""
    if not st.session_state.long_term_memory_enabled:
        print("Long-term memory disabled. Cannot delete historical chats. Resetting current session.")
        if chat_id == st.session_state.current_chat_id:
            create_new_chat_session_in_memory()
            st.rerun()
        return

    hf_token = os.getenv("HF_TOKEN")
    if not hf_token:
        print("HF_TOKEN environment variable not set. Skipping Hugging Face deletion.")
        st.error("Cannot delete chat: Hugging Face token not configured.")
        return

    # Check if the chat to be deleted is the currently active one
    is_current_chat = (chat_id == st.session_state.current_chat_id)

    try:
        # Update in-memory session state first
        if chat_id in st.session_state.all_chat_messages:
            del st.session_state.all_chat_messages[chat_id]
        if chat_id in st.session_state.chat_metadata:
            del st.session_state.chat_metadata[chat_id]
        
        # Save updated metadata and messages to Hugging Face
        # This effectively removes the chat from the JSON files
        save_chat_metadata(st.session_state.user_id, st.session_state.chat_metadata)
        
        # For a full deletion, we need to reload all messages, remove the specific chat_id, and then save the *entire* messages dict.
        messages_filename_in_repo = f"user_memories/{st.session_state.user_id}_messages.json"
        messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"
        
        # Load current messages, remove the specific chat_id, then save the whole thing back
        try:
            existing_messages_content = fs.read_text(messages_hf_path, token=hf_token)
            existing_messages = json.loads(existing_messages_content)
        except FileNotFoundError:
            existing_messages = {}
        
        if chat_id in existing_messages:
            del existing_messages[chat_id]
            with fs.open(messages_hf_path, "w", token=hf_token) as f:
                f.write(json.dumps(existing_messages, indent=2))
            print(f"Chat history for chat {chat_id} explicitly removed from {messages_filename_in_repo} on Hugging Face.")
        else:
            print(f"Chat history for chat {chat_id} not found in {messages_filename_in_repo} on Hugging Face.")


        print(f"Chat '{chat_id}' deleted from in-memory state and updated on Hugging Face.")

        # If the deleted chat was the current one, switch to another or create a new one
        if is_current_chat:
            if st.session_state.chat_metadata:
                # Switch to the first available chat
                first_available_chat_id = next(iter(st.session_state.chat_metadata))
                print(f"Deleted current chat. Switching to: {first_available_chat_id}")
                # Call switch_chat to handle updating session state and rerunning
                switch_chat(first_available_chat_id)
            else:
                # No other chats left, set to a "no chat" state
                print("Deleted last chat. Starting a new empty chat.")
                st.session_state.current_chat_id = None
                st.session_state.messages = [{"role": "assistant", "content": _get_initial_greeting_text()}]
                st.session_state.chat_modified = False
                st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages) # Generate prompts for new empty chat
                st.rerun() # Rerun to display the new state
        else:
            # If a non-current chat was deleted, just rerun to update the sidebar
            st.rerun()
    except Exception as e:
        print(f"Error deleting chat {chat_id} from Hugging Face: {e}")
        st.error(f"Error deleting chat from cloud: {e}")
        # No rerun needed if chat_id wasn't found, as nothing changed.

def rename_chat(chat_id: str, new_name: str): # Modified to accept chat_id
    """Renames the specified chat."""
    if not st.session_state.long_term_memory_enabled:
        print("Long-term memory disabled. Cannot rename chats.")
        return
    if chat_id and new_name and new_name != st.session_state.chat_metadata.get(chat_id):
        st.session_state.chat_metadata[chat_id] = new_name
        save_chat_metadata(st.session_state.user_id, st.session_state.chat_metadata)
        print(f"Renamed chat '{chat_id}' to '{new_name}'")
        # Removed st.rerun() from here as it causes the "no-op" warning in on_change callbacks.
        # Streamlit will automatically rerun after the on_change event completes.

def get_discussion_markdown(chat_id: str) -> str:
    """Retrieves messages for a given chat_id and formats them into a Markdown string."""
    messages = st.session_state.all_chat_messages.get(chat_id, [])
    markdown_content = []
    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"]
        markdown_content.append(f"**{role}:**\n{content}\n\n---")
    return "\n".join(markdown_content)

def get_discussion_docx(chat_id: str) -> bytes:
    """Retrieves messages for a given chat_id and formats them into a DOCX file."""
    messages = st.session_state.all_chat_messages.get(chat_id, [])
    document = Document()
    
    document.add_heading(f"Chat Discussion: {st.session_state.chat_metadata.get(chat_id, 'Untitled Chat')}", level=1)
    # Use current_chat_name if available, otherwise default to a generic name
    document.add_paragraph(f"Exported on: {st.session_state.chat_metadata.get(chat_id, 'Unknown Chat')}") 

    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"]
        
        document.add_heading(f"{role}:", level=3)
        document.add_paragraph(content)
        document.add_paragraph("---") # Separator

    # Save document to a BytesIO object
    byte_stream = BytesIO()
    document.save(byte_stream)
    byte_stream.seek(0) # Rewind to the beginning of the stream
    return byte_stream.getvalue()

def handle_user_input(chat_input_value: str | None):
    """
    Process user input (either from chat box or suggested prompt)
    and update chat with AI response.
    """
    prompt_to_process = None

    if hasattr(st.session_state, 'prompt_to_use') and st.session_state.prompt_to_use:
        prompt_to_process = st.session_state.prompt_to_use
        st.session_state.prompt_to_use = None
    elif chat_input_value:
        prompt_to_process = chat_input_value

    if prompt_to_process:
        # If this is the first user message in a new, unsaved chat, mark it as modified
        # and save its metadata for the first time.
        if not st.session_state.chat_modified and st.session_state.current_chat_id is None:
            new_chat_id = create_new_chat_session_in_memory()
            if st.session_state.long_term_memory_enabled:
                save_chat_metadata(st.session_state.user_id, st.session_state.chat_metadata)
            st.session_state.chat_modified = True
            print(f"Activated new chat '{st.session_state.chat_metadata.get(st.session_state.current_chat_id)}'.")
        elif not st.session_state.chat_modified and len(st.session_state.messages) == 1 and st.session_state.messages[0]["role"] == "assistant":
            st.session_state.chat_modified = True 
            if st.session_state.long_term_memory_enabled:
                save_chat_metadata(st.session_state.user_id, st.session_state.chat_metadata)
            print(f"Chat '{st.session_state.chat_metadata.get(st.session_state.current_chat_id)}' activated and metadata saved.")

        current_user_query = prompt_to_process

        # Format history *before* adding the current user query to st.session_state.messages
        history_for_agent = format_chat_history(st.session_state.messages)

        # Append the user's message to session state so it's displayed by stui.py
        st.session_state.messages.append({"role": "user", "content": current_user_query})

        # Display user message immediately (stui handles this based on session_state.messages)
        # We need to trigger a rerun for stui to pick up the user message before streaming assistant response
        # However, st.write_stream itself will manage UI updates for the streaming part.
        # The critical part is that stui.create_interface is called each time to draw messages.

        # Call get_agent_response with the current query and the prepared history
        response_generator = get_agent_response(current_user_query, chat_history=history_for_agent)

        # Use st.chat_message context manager to ensure assistant's avatar and layout
        # This part is handled by stui.py which iterates through st.session_state.messages.
        # We need to ensure the UI is updated to show the user's message *before* starting the stream.
        # A simple st.rerun() here might be too disruptive if stui is already trying to draw.
        # The goal is that the user message appears, then the assistant message streams in.
        # stui.py will draw all messages in st.session_state.messages.
        # So, we don't add the assistant message to st.session_state.messages yet.
        # Instead, we display it directly using st.write_stream within a chat_message context.

        # The stui.create_interface function will render existing messages.
        # For the new, streaming message, we render it here.
        # This requires that handle_user_input is called in a place where it can write to the main area.

        # This approach assumes handle_user_input is called where it can directly use st.write_stream.
        # If stui.py is the one rendering all messages based on session_state, this needs careful handling.
        # For now, let's assume st.write_stream can be called here.
        # We will append to messages *after* the stream is complete.

        # No, the stui.py renders messages. We should not use st.chat_message here.
        # We need to append a placeholder to messages, let stui render it, and then update it.
        # However, st.write_stream is designed to be used directly.

        # Let's adjust: st.write_stream will be called by stui.py or a similar UI rendering function.
        # This function (handle_user_input) should prepare the generator and potentially store it.
        # This is getting complicated. The instruction says "use st.write_stream to display".

        # Re-evaluating: The `stui.create_interface` is likely called once.
        # The chat input is handled by `handle_user_input_callback` which points to this function.
        # So, this function IS where UI interactions happen for the chat.

        # Add a temporary placeholder for the assistant message for stui.py to render
        # This is a bit of a hack if st.write_stream is to be used directly *below*.
        # Let's follow the prompt: use st.write_stream directly for display.
        # This means that the display of this specific message will happen outside the stui.py loop for previous messages.

        # The user message is already added to st.session_state.messages.
        # stui.py will render all messages in st.session_state.messages *up to the last one*.
        # Then we stream the new one.

        full_response_content = ""
        # The user message is in st.session_state.messages.
        # The stui.create_interface will draw all messages from st.session_state.messages.
        # So, before we call st.write_stream, all *previous* messages should be on screen.

        # We need to make sure handle_user_input is called in a context where it can write to the chat display area.
        # Assuming stui.create_interface sets up the general layout, and then the input handling
        # can add new messages to the chat area.

        # This is where the new assistant message will appear
        # with st.chat_message("assistant"): # This is usually handled by the UI loop (stui.py)
            # Let's assume stui.py has already rendered previous messages.
            # Now we stream the new one.
            # full_response_content = st.write_stream(response_generator) # This will render directly.

        # The prompt implies handle_user_input itself should call st.write_stream.
        # This means stui.py might need to be structured to allow this.
        # For now, let's assume it's okay to call st.write_stream here.
        # The stui.py will render all messages based on st.session_state.messages on each run.
        # So, we should append the user message, then immediately prepare for assistant's response.

        # The challenge: st.write_stream must be called in the main script execution flow, not inside a callback
        # if the callback doesn't directly manipulate Streamlit elements.
        # However, `handle_user_input` *is* the callback.
        # Let's assume `stui.create_interface` calls this, and then this function uses Streamlit commands.

        # The previous messages are rendered by stui.py.
        # Now, for the new response:
        # This will write the stream to wherever the current Streamlit cursor is.
        # This assumes `stui.py` has set up the chat interface area.
        # We will append the user message, then the assistant response will be streamed.
        # After streaming, we append the full assistant response to messages for history.

        # The stui.py will call this. It will display all messages from st.session_state.messages.
        # Then this function will stream the new response.
        # This seems like the correct flow.

        # We don't need a placeholder in st.session_state.messages for st.write_stream.
        # st.write_stream handles its own UI output.
        with st.chat_message("assistant"): # To get the avatar and proper styling.
            full_response_content = st.write_stream(response_generator)

        st.session_state.messages.append({"role": "assistant", "content": full_response_content})

        # Autosave the current chat history after AI response if it's been modified
        if st.session_state.chat_modified and st.session_state.long_term_memory_enabled:
            save_chat_history(st.session_state.user_id, st.session_state.current_chat_id, st.session_state.messages)

        st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages) # Use cached version
        st.rerun()

def reset_chat_callback():
    """Resets the chat by creating a new, unsaved chat session."""
    print("Resetting chat by creating a new session...")
    create_new_chat_session_in_memory() # Create new chat in memory
    st.rerun() # Rerun to display the new chat

def handle_regeneration_request():
    """Handles the request to regenerate the last assistant response."""
    if not st.session_state.get("do_regenerate", False):
        return

    st.session_state.do_regenerate = False

    if not st.session_state.messages or st.session_state.messages[-1]['role'] != 'assistant':
        print("Warning: Regeneration called but last message is not from assistant or no messages exist.")
        st.rerun()
        return

    if len(st.session_state.messages) == 1: # Initial greeting
        print("Regenerating initial greeting (simulated stream)...")
        # For the initial greeting, we stream it directly and update the content.
        response_generator = get_agent_response("Regenerate initial greeting", []) # Query content doesn't really matter here

        # The stui.py will render based on st.session_state.messages.
        # We need to update messages[0] after streaming.
        with st.chat_message("assistant"): # This is for display during regeneration
            full_response_content = st.write_stream(response_generator)

        st.session_state.messages[0]['content'] = full_response_content # Update existing greeting

        if st.session_state.long_term_memory_enabled:
            save_chat_history(st.session_state.user_id, st.session_state.current_chat_id, st.session_state.messages)
        st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages)
        st.rerun()
        return

    print("Regenerating last assistant response (simulated stream)...")
    # Remove last assistant message from state, the UI will reflect this on next rerun (implicitly via st.write_stream)
    st.session_state.messages.pop()

    if not st.session_state.messages or st.session_state.messages[-1]['role'] != 'user':
        print("Warning: Cannot regenerate, no preceding user query found after popping assistant message.")
        # If we can't regenerate, add back a generic message or handle error
        st.session_state.messages.append({"role": "assistant", "content": "Could not regenerate the previous response."})
        st.rerun() # Rerun to show the error/fallback message
        return

    prompt_to_regenerate = st.session_state.messages[-1]['content']
    # History for regeneration should be up to the message *before* the user query that led to the response being regenerated.
    # So, it's all messages *excluding* the last user message.
    history_for_regen = format_chat_history(st.session_state.messages[:-1]) # Exclude current user prompt

    response_generator = get_agent_response(prompt_to_regenerate, chat_history=history_for_regen)

    # Display the streaming response using st.write_stream in the context of an assistant message
    # This assumes stui.py has already rendered messages up to the point of the one being replaced.
    with st.chat_message("assistant"):
        full_response_content = st.write_stream(response_generator)

    # Now, add this regenerated response as a new message.
    # Or, if stui.py expects the last message to be updated:
    st.session_state.messages.append({"role": "assistant", "content": full_response_content})
    # This will mean the regenerated response appears as a new message.
    # If the goal is to *replace* the popped message, the logic in stui.py would need to handle that,
    # or we would update st.session_state.messages[-1]['content'] if a placeholder was added.
    # Given we popped, then do a st.write_stream, then append, this is effectively replacing by adding anew.

    if st.session_state.long_term_memory_enabled:
        save_chat_history(st.session_state.user_id, st.session_state.current_chat_id, st.session_state.messages)
    st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages)
    st.rerun()

def forget_me_and_reset():
    """
    Deletes all user chat histories from the Hugging Face JSON files, removes the user ID cookie,
    and resets the Streamlit session state to a fresh start.
    """
    user_id_to_delete = st.session_state.get("user_id")
    hf_token = os.getenv("HF_TOKEN")

    if user_id_to_delete and hf_token:
        try:
            # Use HF_USER_MEMORIES_DATASET_ID for user memories
            metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id_to_delete}_metadata.json"
            messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id_to_delete}_messages.json"

            # Attempt to delete both files
            deleted_any = False
            try:
                fs.rm(metadata_hf_path, token=hf_token)
                print(f"Deleted metadata file for user '{user_id_to_delete}' from Hugging Face.")
                deleted_any = True
            except FileNotFoundError:
                print(f"Metadata file for user '{user_id_to_delete}' not found, skipping deletion.")
            except Exception as e:
                print(f"Error deleting metadata file for user '{user_id_to_delete}': {e}")
                st.error(f"Failed to delete metadata from cloud: {e}")

            try:
                fs.rm(messages_hf_path, token=hf_token)
                print(f"Deleted messages file for user '{user_id_to_delete}' from Hugging Face.")
                deleted_any = True
            except FileNotFoundError:
                print(f"Messages file for user '{user_id_to_delete}' not found, skipping deletion.")
            except Exception as e:
                print(f"Error deleting messages file for user '{user_id_to_delete}': {e}")
                st.error(f"Failed to delete messages from cloud: {e}")

            if deleted_any:
                print(f"Successfully attempted to delete all data for user '{user_id_to_delete}' from Hugging Face.")
            else:
                print(f"No files found to delete for user '{user_id_to_delete}' on Hugging Face.")

        except Exception as e:
            print(f"General error during Hugging Face deletion for user {user_id_to_delete}: {e}")
            st.error(f"Failed to delete user data from cloud: {e}")
    elif not hf_token:
        print("HF_TOKEN environment variable not set. Cannot delete user data from Hugging Face.")
        st.warning("Cannot delete user data from cloud: Hugging Face token not configured.")

    # Delete the user ID cookie
    try:
        cookies.delete(cookie="user_id")
        print(f"Deleted user ID cookie for '{user_id_to_delete}'")
    except Exception as e:
        print(f"ERROR: Failed to delete user_id cookie for {user_id_to_delete}: {e}")
        st.error(f"Failed to delete user ID cookie: {e}")

    # Reset session state to clear all chat history and user data in memory
    st.session_state.chat_metadata = {}
    st.session_state.all_chat_messages = {}
    st.session_state.current_chat_id = None
    st.session_state.messages = [] # Clear messages, will be re-populated by main's init
    st.session_state.chat_modified = False
    st.session_state.suggested_prompts = DEFAULT_PROMPTS # Reset to default prompts
    st.session_state.renaming_chat_id = None
    st.session_state.uploaded_documents = {}
    st.session_state.uploaded_dataframes = {}
    
    # Crucially, reset the session_control_flags_initialized to force full re-initialization
    # in the main function on the next rerun.
    st.session_state.session_control_flags_initialized = False
    st.session_state._greeting_logic_log_shown_for_current_state = False # Reset for fresh log on next run
    
    # Delete user_id from session state to force re-generation of a temporary one
    if "user_id" in st.session_state:
        del st.session_state.user_id

    # Use JavaScript to clear cookies and force a full page reload
    # This ensures a complete reset from the browser's perspective.
    js_code = """
    <script>
        function deleteAllCookies() {
            const cookies = document.cookie.split(';');
            for (let i = 0; i < cookies.length; i++) {
                const cookie = cookies[i];
                const eqPos = cookie.indexOf('=');
                const name = eqPos > -1 ? cookie.substr(0, eqPos) : cookie;
                document.cookie = name + '=;expires=Thu, 01 Jan 1970 00:00:00 GMT;path=/';
            }
        }
        deleteAllCookies();
        window.location.reload(true); // Force a hard reload from the server
    </script>
    """
    st.components.v1.html(js_code, height=0, width=0)

    print(f"Session reset. New temporary user ID will be generated on next run.")
    # No st.rerun() here, as the JavaScript reload will handle it.

def _set_long_term_memory_preference():
    """Callback to save the long_term_memory_enabled state to a cookie."""
    current_value = st.session_state.long_term_memory_enabled
    try:
        cookies.set(cookie="long_term_memory_pref", val=str(current_value))
        print(f"Long-term memory preference saved to cookie: {current_value}")
    except Exception as e:
        print(f"ERROR: Failed to save long-term memory preference to cookie: {e}")
        st.error(f"Failed to save preference: {e}")
    # No st.rerun() here, as the toggle itself triggers a rerun.
    # The main loop's memory state change detection will handle the rest.
    st.session_state._last_memory_state_changed_by_toggle = True

def main():
    """Main function to run the Streamlit app."""
    success, error_message = setup_global_llm_settings()
    if not success:
        st.error(error_message)
        st.stop()

    # --- Long-term memory initialization and change detection ---
    pref_from_cookie = cookies.get(cookie="long_term_memory_pref")

    if "long_term_memory_enabled" not in st.session_state:
        if pref_from_cookie is not None:
            pref_str = str(pref_from_cookie).lower()
            
            if pref_str == 'true' or pref_str == '1':
                st.session_state.long_term_memory_enabled = True
            elif pref_str == 'false' or pref_str == '0':
                st.session_state.long_term_memory_enabled = False
            else:
                st.session_state.long_term_memory_enabled = True
                print(f"Warning: Unexpected value for long_term_memory_pref cookie: '{pref_from_cookie}'. Defaulting to True.")
            print(f"Long-term memory preference loaded from cookie: {st.session_state.long_term_memory_enabled}")
        else:
            st.session_state.long_term_memory_enabled = True  # Default: enabled
            cookies.set(cookie="long_term_memory_pref", val=str(st.session_state.long_term_memory_enabled))
            print(f"Long-term memory preference not found. Defaulting to {st.session_state.long_term_memory_enabled} and saving cookie.")

    if "_last_memory_state_was_enabled" not in st.session_state:
        st.session_state._last_memory_state_was_enabled = st.session_state.long_term_memory_enabled

    # --- Handle Memory State Change ---
    memory_state_has_changed_this_run = st.session_state._last_memory_state_was_enabled != st.session_state.long_term_memory_enabled
    if memory_state_has_changed_this_run:
        print(f"Memory state changed from {st.session_state._last_memory_state_was_enabled} to {st.session_state.long_term_memory_enabled}. Re-initializing session.")
        st.session_state._last_memory_state_was_enabled = st.session_state.long_term_memory_enabled
        st.session_state.session_control_flags_initialized = False

        if "user_id" in st.session_state:
            del st.session_state.user_id
        
        _initialize_user_session_data.clear()
        print("Cleared user data cache due to memory state change.")

    if "_last_memory_state_changed_by_toggle" not in st.session_state:
        st.session_state._last_memory_state_changed_by_toggle = False

    st.session_state._last_memory_state_changed_by_toggle = False


    # --- Core Session Variable Initialization (runs once per session OR after memory state change) ---
    if not st.session_state.get("session_control_flags_initialized", False):
        print("Initializing core session variables...")

        st.session_state.initial_greeting_shown_for_session = False
        st.session_state.current_chat_id = None
        st.session_state.messages = []
        st.session_state.chat_modified = False
        st.session_state.suggested_prompts = DEFAULT_PROMPTS
        st.session_state.renaming_chat_id = None
        st.session_state.uploaded_documents = {}
        st.session_state.uploaded_dataframes = {}

        st.session_state.session_control_flags_initialized = True
        print("Core session variables initialized.")

    # --- User ID and Chat Data Load (cached, sensitive to memory state) ---
    user_id_val, chat_metadata_val, all_chat_messages_val, cookie_action = \
        _initialize_user_session_data(st.session_state.long_term_memory_enabled)

    st.session_state.user_id = user_id_val
    st.session_state.chat_metadata = chat_metadata_val
    st.session_state.all_chat_messages = all_chat_messages_val

    # --- Apply cookie actions based on _initialize_user_session_data result ---
    if cookie_action == "SET_COOKIE":
        import datetime
        expires = datetime.datetime.now() + datetime.timedelta(days=365)
        cookies.set(cookie="user_id", val=user_id_val, expires_at=expires)
        print(f"Set user_id cookie: {user_id_val}")
    elif cookie_action == "DELETE_COOKIE":
        cookies.delete(cookie="user_id")
        print("Deleted user_id cookie.")

    # --- Agent Initialization (runs once per session) ---
    if AGENT_SESSION_KEY not in st.session_state:
        agent_instance, error_message = setup_agent(max_search_results=10) 
        if agent_instance is None:
            st.error(error_message)
            st.stop()
        st.session_state[AGENT_SESSION_KEY] = agent_instance

    # --- Active Chat and Initial Greeting Logic ---
    should_rerun_after_init = False

    if st.session_state.long_term_memory_enabled:
        if st.session_state.current_chat_id and st.session_state.current_chat_id in st.session_state.chat_metadata:
            if st.session_state.all_chat_messages.get(st.session_state.current_chat_id) is None:
                print(f"WARNING: Messages for current chat ID '{st.session_state.current_chat_id}' were not loaded. Setting to empty list.")
                st.session_state.all_chat_messages[st.session_state.current_chat_id] = []
            
            st.session_state.messages = st.session_state.all_chat_messages.get(st.session_state.current_chat_id, [])
            st.session_state.chat_modified = True
            # print(f"Active chat is '{st.session_state.current_chat_id}'. Messages: {len(st.session_state.messages)}") # Removed verbose log

        elif st.session_state.chat_metadata:
            first_available_chat_id = next(iter(st.session_state.chat_metadata))
            print(f"No current chat ID. Selecting first available: '{first_available_chat_id}'.")
            st.session_state.current_chat_id = first_available_chat_id
            st.session_state.messages = st.session_state.all_chat_messages.get(first_available_chat_id, [])
            st.session_state.chat_modified = True
            should_rerun_after_init = True
        else:
            if not st.session_state.initial_greeting_shown_for_session:
                print("No chats exist. Displaying initial greeting.")
                st.session_state.messages = [{"role": "assistant", "content": _get_initial_greeting_text()}]
                st.session_state.initial_greeting_shown_for_session = True
                st.session_state.current_chat_id = None
                st.session_state.chat_modified = False
                should_rerun_after_init = True
    else:
        if not st.session_state.current_chat_id or \
           st.session_state.current_chat_id not in st.session_state.all_chat_messages or \
           not st.session_state.messages:
            if not st.session_state.initial_greeting_shown_for_session:
                print("Creating new temporary session with greeting.")
                create_new_chat_session_in_memory()
                st.session_state.initial_greeting_shown_for_session = True
                should_rerun_after_init = True
            elif not st.session_state.messages:
                 print("Messages empty, recreating greeting for temporary session.")
                 create_new_chat_session_in_memory()
                 should_rerun_after_init = True
            else:
                 st.session_state.messages = st.session_state.all_chat_messages[st.session_state.current_chat_id]
                 # print(f"Using existing temporary chat '{st.session_state.current_chat_id}'. Messages: {len(st.session_state.messages)}") # Removed verbose log
        else:
            st.session_state.messages = st.session_state.all_chat_messages[st.session_state.current_chat_id]
            # print(f"Confirmed existing temporary chat '{st.session_state.current_chat_id}'. Messages: {len(st.session_state.messages)}") # Removed verbose log

    # Fallback: If messages list is somehow still not a list (should be extremely rare)
    if not isinstance(st.session_state.messages, list):
        print("WARNING: Session messages not a list. Resetting to empty list and default prompts.")
        st.session_state.messages = []
        st.session_state.suggested_prompts = DEFAULT_PROMPTS
        st.session_state.current_chat_id = None
        st.session_state.chat_modified = False
        should_rerun_after_init = True
    elif not st.session_state.messages and not st.session_state.initial_greeting_shown_for_session:
        print("FALLBACK: No messages and no greeting shown. Displaying initial greeting.")
        st.session_state.messages = [{"role": "assistant", "content": _get_initial_greeting_text()}]
        st.session_state.initial_greeting_shown_for_session = True
        st.session_state.current_chat_id = None
        st.session_state.chat_modified = False
        should_rerun_after_init = True

    # Update suggested prompts based on the final state of messages
    if 'suggested_prompts' not in st.session_state or \
       (st.session_state.messages and st.session_state.suggested_prompts == DEFAULT_PROMPTS and len(st.session_state.messages) > 1) or \
       (not st.session_state.messages and st.session_state.suggested_prompts != DEFAULT_PROMPTS):
        print("Updating suggested prompts.")
        st.session_state.suggested_prompts = _cached_generate_suggested_prompts(st.session_state.messages if st.session_state.messages else [])
        if not should_rerun_after_init and len(st.session_state.messages) == 1 and st.session_state.messages[0]["role"] == "assistant":
             should_rerun_after_init = True


    # Final check for rerun after initial chat setup
    if should_rerun_after_init:
        st.rerun()

    if st.session_state.get("do_regenerate", False):
        handle_regeneration_request()

    stui.create_interface(
        reset_callback=reset_chat_callback,
        new_chat_callback=lambda: create_new_chat_session_in_memory() and st.rerun(),
        delete_chat_callback=delete_chat_session,
        rename_chat_callback=rename_chat, # Pass the modified rename_chat function
        chat_metadata=st.session_state.chat_metadata,
        current_chat_id=st.session_state.current_chat_id,
        switch_chat_callback=switch_chat,
        get_discussion_markdown_callback=get_discussion_markdown,
        get_discussion_docx_callback=get_discussion_docx, # Pass the new DOCX callback
        suggested_prompts_list=st.session_state.suggested_prompts,
        handle_user_input_callback=handle_user_input,
        long_term_memory_enabled=st.session_state.long_term_memory_enabled, # Pass the new setting
        forget_me_callback=forget_me_and_reset, # Pass the new callback
        set_long_term_memory_callback=_set_long_term_memory_preference # Pass the new callback
    )

    chat_input_for_handler = st.session_state.get("chat_input_value_from_stui")
    if "chat_input_value_from_stui" in st.session_state: # Ensure the key exists before deleting
        del st.session_state.chat_input_value_from_stui # Or set to None: st.session_state.chat_input_value_from_stui = None
    
    # Call handle_user_input if there's a chat input or a suggested prompt pending
    if chat_input_for_handler or st.session_state.get('prompt_to_use'):
        handle_user_input(chat_input_for_handler)

if __name__ == "__main__":
    if not os.getenv("GOOGLE_API_KEY"):
        st.warning("⚠️ GOOGLE_API_KEY environment variable not set. The agent may not work properly.")
    main()
